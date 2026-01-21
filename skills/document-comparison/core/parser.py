from pathlib import Path
from typing import Dict, Any
from docx import Document
import re
import zipfile
from lxml import etree

class DocumentParser:
    """Parse Word documents into structured data."""

    def parse_document(self, filepath: str) -> Dict[str, Any]:
        """Parse a Word document and return structured data."""
        if not Path(filepath).exists():
            raise FileNotFoundError(f"Document not found: {filepath}")

        doc = Document(filepath)

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text,
                    'style': para.style.name,
                    'level': self._get_heading_level(para.style.name)
                })

        tables = []
        for table in doc.tables:
            table_data = {
                'rows': len(table.rows),
                'cols': len(table.columns),
                'cells': []
            }

            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data['cells'].append(row_data)

            tables.append(table_data)

        sdt_paragraphs, sdt_tables = self._parse_sdt_elements(filepath)

        all_paragraphs = paragraphs + sdt_paragraphs
        all_tables = tables + sdt_tables

        result = {
            'paragraphs': all_paragraphs,
            'tables': all_tables,
            'metadata': {
                'paragraph_count': len(all_paragraphs),
                'table_count': len(all_tables),
                'standard_tables': len(tables),
                'sdt_tables': len(sdt_tables),
                'sdt_paragraphs': len(sdt_paragraphs)
            }
        }

        return result

    def _parse_sdt_elements(self, filepath: str) -> tuple[list[Dict], list[Dict]]:
        """Parse SDT (Structured Document Tag) elements for enhanced table extraction."""
        sdt_paragraphs = []
        sdt_tables = []

        try:
            with zipfile.ZipFile(filepath, 'r') as zip_ref:
                xml_content = zip_ref.read('word/document.xml')
                xml_str = xml_content.decode('utf-8')

                tbl_pattern = r'<w:tbl[^>]*>(.*?)</w:tbl>'
                all_tables = re.findall(tbl_pattern, xml_str, re.DOTALL)

                for tbl in all_tables:
                    tr_pattern = r'<w:tr[^>]*>(.*?)</w:tr>'
                    rows = re.findall(tr_pattern, tbl, re.DOTALL)

                    table_data = {
                        'rows': len(rows),
                        'cols': 0,
                        'cells': []
                    }

                    for row in rows:
                        tc_pattern = r'<w:t[^>]*>([^<]*)</w:t>'
                        texts = re.findall(tc_pattern, row)
                        cell_text = ''.join(texts).strip()
                        table_data['cells'].append([cell_text])

                    if len(table_data['cells']) > 0:
                        table_data['cols'] = max(len(row) for row in table_data['cells'])
                        sdt_tables.append(table_data)

                xml_no_tables = re.sub(r'<w:tbl[^>]*>.*?</w:tbl>', '', xml_str, flags=re.DOTALL)

                p_pattern = r'<w:t[^>]*>([^<]*)</w:t>'
                all_texts = re.findall(p_pattern, xml_no_tables)

                current_text = ''
                for text in all_texts:
                    if text.strip():
                        current_text += text
                    else:
                        if current_text.strip():
                            sdt_paragraphs.append({
                                'text': current_text.strip(),
                                'style': 'Normal',
                                'level': 0
                            })
                        current_text = ''

                if current_text.strip():
                    sdt_paragraphs.append({
                        'text': current_text.strip(),
                        'style': 'Normal',
                        'level': 0
                    })

        except Exception as e:
            print(f"Warning: Failed to parse SDT elements: {e}")

        return sdt_paragraphs, sdt_tables

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        if 'Heading 1' in style_name:
            return 1
        elif 'Heading 2' in style_name:
            return 2
        elif 'Heading 3' in style_name:
            return 3
        return 0
